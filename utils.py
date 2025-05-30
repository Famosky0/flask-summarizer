import fitz  # PyMuPDF

def extract_text_from_pdf(file_path):
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

# utils.py
from docx import Document

def export_summary_to_docx(summaries: dict, output_path: str = "summary.docx"):
    doc = Document()
    doc.add_heading("Summarized Academic Paper", 0)
    for section_title, summary in summaries.items():
        doc.add_heading(section_title, level=1)
        doc.add_paragraph(summary)
    doc.save(output_path)

def export_references_to_docx(references_text: str, output_path: str = "references.docx"):
    doc = Document()
    doc.add_heading("References", 0)
    doc.add_paragraph(references_text)
    doc.save(output_path)
